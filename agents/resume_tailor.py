"""
agents/resume_tailor.py
───────────────────────
Uses Claude AI to tailor the base resume to each job description.
Modifies only the summary, skills, and experience bullet points.
Outputs a tailored .docx file per job.

Environment vars needed:
  ANTHROPIC_API_KEY   — Anthropic API key
"""

import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import anthropic
import yaml

logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
with open(ROOT / "config" / "settings.yaml") as f:
    SETTINGS = yaml.safe_load(f)

BASE_RESUME = ROOT / SETTINGS["resume"]["base_resume_path"]
TAILORED_DIR = ROOT / SETTINGS["resume"]["tailored_output_dir"]
MODEL = SETTINGS["resume"].get("ai_model", "claude-sonnet-4-20250514")

TAILORED_DIR.mkdir(parents=True, exist_ok=True)


# ── Resume text extraction ────────────────────────────────────────────────────

def extract_resume_text(docx_path: Path) -> str:
    """Extract plain text from .docx using python-docx."""
    try:
        from docx import Document
        doc = Document(str(docx_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.error(f"Failed to extract resume text: {e}")
        return ""


# ── AI tailoring ─────────────────────────────────────────────────────────────

TAILOR_SYSTEM_PROMPT = """You are an expert resume writer specializing in DevOps, Cloud, and Automation roles.
Your task is to tailor a candidate's resume to closely match a job description.

Rules:
1. ONLY modify: Professional Summary, Skills section, and bullet points under Experience.
2. NEVER change: candidate name, contact info, education, certifications, company names, job titles, or dates.
3. Mirror the language and keywords from the JD naturally — don't keyword-stuff.
4. Keep bullet points concise, action-verb-led, and quantified where possible.
5. Return ONLY the tailored resume text, preserving the same section structure.
6. Do not add sections that don't exist in the original resume.
7. Output format must be valid for re-insertion into a Word document.
"""

def tailor_resume_with_ai(resume_text: str, job: dict) -> str:
    """Call Claude to tailor the resume text to the job description."""
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    user_prompt = f"""
## Job Title: {job['title']}
## Company: {job['company']}
## Location: {job['location']}

## Job Description:
{job['description'][:4000]}  

## Current Resume:
{resume_text}

Please tailor the resume to this job description following the rules above.
"""

    message = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        system=TAILOR_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
    )

    return message.content[0].text


# ── DOCX writing ─────────────────────────────────────────────────────────────

def write_tailored_docx(tailored_text: str, job: dict, output_path: Path):
    """
    Write the tailored text into a copy of the base resume .docx.
    Strategy: copy the base .docx, then replace paragraph text using python-docx.
    This preserves formatting while updating content.
    """
    from docx import Document

    shutil.copy2(str(BASE_RESUME), str(output_path))
    doc = Document(str(output_path))

    tailored_lines = [l.strip() for l in tailored_text.splitlines() if l.strip()]
    original_lines = [p.text.strip() for p in doc.paragraphs]

    # Build a mapping of paragraph index → new text
    # (simple line-by-line replacement; assumes same structure)
    tailored_idx = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if tailored_idx < len(tailored_lines):
            # Preserve run formatting; update the first run's text
            if para.runs:
                para.runs[0].text = tailored_lines[tailored_idx]
                # Clear remaining runs in this paragraph to avoid duplication
                for run in para.runs[1:]:
                    run.text = ""
            tailored_idx += 1

    doc.save(str(output_path))
    logger.info(f"💾 Tailored resume saved: {output_path.name}")


# ── Main class ────────────────────────────────────────────────────────────────

class ResumeTailor:
    def __init__(self):
        if not BASE_RESUME.exists():
            raise FileNotFoundError(
                f"Base resume not found at {BASE_RESUME}. "
                "Please place your resume at data/resumes/base_resume.docx"
            )

    def tailor_for_job(self, job: dict) -> Path | None:
        """
        Tailor the resume for a single job. Returns path to tailored .docx or None on failure.
        """
        job_id = job["id"]
        safe_company = re.sub(r"[^\w\-]", "_", job.get("company", "Company"))
        safe_title = re.sub(r"[^\w\-]", "_", job.get("title", "Role"))[:40]
        output_name = f"{safe_title}_{safe_company}_{job_id[:8]}.docx"
        output_path = TAILORED_DIR / output_name

        if output_path.exists():
            logger.info(f"Tailored resume already exists for job {job_id}: {output_name}")
            return output_path

        logger.info(f"✍️  Tailoring resume for: {job['title']} @ {job['company']}")

        resume_text = extract_resume_text(BASE_RESUME)
        if not resume_text:
            logger.error("Could not extract resume text. Skipping.")
            return None

        try:
            tailored_text = tailor_resume_with_ai(resume_text, job)
            write_tailored_docx(tailored_text, job, output_path)
            return output_path
        except Exception as e:
            logger.error(f"Failed to tailor resume for {job_id}: {e}")
            return None

    def tailor_batch(self, jobs: list[dict]) -> dict[str, Path]:
        """Tailor resumes for a list of jobs. Returns {job_id: resume_path}."""
        results = {}
        for job in jobs:
            path = self.tailor_for_job(job)
            if path:
                results[job["id"]] = path
        return results


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    # Quick test with a dummy job
    tailor = ResumeTailor()
    test_job = {
        "id": "test-001",
        "title": "Senior DevOps Engineer",
        "company": "Acme Corp",
        "location": "Dubai, UAE",
        "description": "We need a DevOps engineer with Kubernetes, Ansible, Python, Terraform, CI/CD experience.",
    }
    path = tailor.tailor_for_job(test_job)
    print(f"Tailored resume: {path}")
